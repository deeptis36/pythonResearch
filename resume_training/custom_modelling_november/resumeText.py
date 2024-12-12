import re
import fitz  # PyMuPDF
import docx
import os
import PyPDF2

from xml.etree import ElementTree
from lxml import etree


def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                text += page.get_text()  # Accumulate text from all pages
    except Exception as e:
        print(f"Error reading PDF: {e}")
    
    return text.strip()  # Ret


def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = docx.Document(docx_path)
   
    text_parts = []
    text = ""
    
    # Extract text from headers
    for section in doc.sections:
        header = section.header
        for para in header.paragraphs:
            text += para.text + "\n"
    
    # Extract text from paragraphs
    for para in doc.paragraphs:
        text += para.text + "\n"

    # Extracting XML to handle shapes (e.g., textboxes) with lxml
    doc_xml = doc._element
    xml_str = etree.tostring(doc_xml)  # Convert the XML tree to a string
    
    # Parse the XML string using lxml to enable XPath queries
    root = etree.fromstring(xml_str)
    
    # Extract text from shapes (textboxes)
    for shape in root.xpath('.//v:shape', namespaces={'v': 'urn:schemas-microsoft-com:vml'}):
        if 'textbox' in shape.attrib.get('{urn:schemas-microsoft-com:vml}type', ''):
            for text_box in shape.findall('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                text_parts.append(text_box.text.strip())
    
    # Combine all extracted text parts
    text += '\n'.join(text_parts)
    
    # Extract text from tables
    if len(doc.tables) > 0:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
   
    return text.strip()


def get_resume_text(file_path):
    """Extract resume text from a PDF or DOCX file."""
    if file_path.endswith('.pdf'):
        resume_text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        resume_text = extract_text_from_docx(file_path)
    else:
        return None  # Return None if the file is neither PDF nor DOCX
    return resume_text.strip()




def get_text_or_delete_if_not_readable(file_path):
    try:
        with open(file_path, 'r') as file:
            # Read and return the file content
            content = file.read()
            print(f"The file '{file_path}' is readable.")
            return content
    except Exception as e:
        # If there's an error reading the file, delete it
        print(f"The file '{file_path}' is not readable: {e}")
        try:
            os.remove(file_path)
            # print(f"The file '{file_path}' has been deleted.")
            return None
        except Exception as delete_error:
            print(f"Failed to delete the file '{file_path}': {delete_error}")
        return None  # Return
    return None