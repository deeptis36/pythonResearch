import re
import fitz  # PyMuPDF
import docx
import os



def extract_text_from_pdf(pdf_path):
    """Extract text from a PDF file using PyMuPDF."""
    doc = fitz.open(pdf_path)
    text = ""
    for page in doc:
        text += page.get_text()
    return text

def extract_text_from_docx(docx_path):
    """Extract text from a DOCX file."""
    doc = docx.Document(docx_path)
   

    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"

    # Checking for tables and extracting text from them
    if len(doc.tables) > 0:
        
        for i, table in enumerate(doc.tables):
           
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + "\n"
    
    return text


def get_resume_text(file_path):
    """Extract resume text from a PDF or DOCX file."""
    if file_path.endswith('.pdf'):
        resume_text = extract_text_from_pdf(file_path)
    elif file_path.endswith('.docx'):
        resume_text = extract_text_from_docx(file_path)
    else:
        return None  # Return None if the file is neither PDF nor DOCX
    return resume_text